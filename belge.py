from docx import Document
import requests
from datetime import datetime
import datetime
import time
import sys
import os
import pandas as pd
import sys
import subprocess
import re
from sys import argv

class ParagraphsKeyWordsReplace:
    def __init__(self):
        self.text = None
        self.runs = None

    def p_replace(self,x,key,value):
        paragraph_positions = []
        for y, run in enumerate(self.runs):
            for z, char in enumerate(list(run.text)):
                position = {"run": y, "char": z}     
                paragraph_positions.append(position)
        key_indexs = [s for s in range(len(self.text)) if self.text.find(key, s, len(self.text)) == s]
        for i, start_i in enumerate( reversed(key_indexs),start=1):
            end_i = start_i + len(key)
            key_maps = paragraph_positions[start_i:end_i]
            ParagraphsKeyWordsReplace.c_replace(self, key_maps, value)


    def c_replace(self,key_maps,value):
        for i, position in enumerate(reversed(key_maps),start=1):
            y, z = position["run"], position["char"]
            run,char = self.runs[y],self.runs[y].text[z]
            if i < len(key_maps):
                rt = list(run.text)
                rt.pop(z)
                run.text = ''.join(rt)
            if i == len(key_maps):
                run.text = run.text.replace(char, value)

class DocxKeyWordsReplace:
    def __init__(self):
        self.paragraphs = None
        self.sections = None

    def content(self,replace_dict):
        for key, value in replace_dict.items():
            for x, paragraph in enumerate(self.paragraphs):
                ParagraphsKeyWordsReplace.p_replace(paragraph,x,key,value)

def personalinfos(login):
    gsheetid = "xxxx"
    sheet_name = ""
    gsheet_url = "https://docs.google.com/spreadsheets/d/{}/gviz/tq?tqx=out:csv&sheet={}".format(gsheetid, sheet_name)
    istanbuldf = pd.read_csv(gsheet_url)

    gsheetid = "xxxx"
    sheet_name = ""
    gsheet_url = "https://docs.google.com/spreadsheets/d/{}/gviz/tq?tqx=out:csv&sheet={}".format(gsheetid, sheet_name)
    kocaelidf = pd.read_csv(gsheet_url)
    searchistanbul = (istanbuldf.loc[istanbuldf["Login"] == login])
    searchkocaeli = (kocaelidf.loc[kocaelidf["Login"] == login])
    istanbullist = searchistanbul.values.tolist()
    kocaelilist = searchkocaeli.values.tolist()
    infolist = []
    if (len(istanbullist) > 0):
        converdatetime = datetime.datetime.strptime(istanbullist[0][2], "%Y-%m-%d")
        convertformat = converdatetime.strftime('%d.%m.%Y')
        infolist.append(convertformat)
        infolist.append(istanbullist[0][3])
        return infolist
    if (len(kocaelilist) > 0):
        converdatetime = datetime.datetime.strptime(kocaelilist[0][2], "%Y-%m-%d")
        convertformat = converdatetime.strftime('%d.%m.%Y')
        infolist.append(convertformat)
        infolist.append(kocaelilist[0][3])
        return infolist
    else:
        return "ahmethakangunessds24@gmail.com"

def getinfo(login, token):
    headers = {
    'Authorization': 'Bearer ' + token,
    }
    endpoint = "/v2/users/{}".format(login)
    response = requests.get('https://api.intra.42.fr' + "{}".format(endpoint), headers=headers)
    if (response.status_code == 200):
        responsejs = response.json()
        list = []
        list.append(responsejs['id'])
        list.append(responsejs['first_name'])
        list.append(responsejs['last_name'])
        list.append(responsejs['login'])
        list.append(datetime.datetime.today().strftime("%d.%m.%Y"))
        return (list)
    else:
        exit (1)

def belge(login, token, slackmail):
    personal = personalinfos(login)
    if (slackmail != personal[1]):
        return (0)
    list = getinfo(login, token)
    replace_dict = {
        "Davut Uzun": str(list[1].title()) + " " + str(list[2].title()),
        "DAVUT": str(list[1]).title(),
        "UZUN": str(list[2].title()),
        "12.08.1973": str(personal[0]),
        "100491": str(list[0]),
        "31.10.2022": list[4]
        }
    docx = Document("belge/student.docx")
    DocxKeyWordsReplace.content(docx, replace_dict=replace_dict)
    docx.save("belge/" + list[3] + ".docx")
    os.system("unoconv -f pdf " + "belge/" + str(login) + ".docx")
    os.system("rm -rf belge/" + str(login) + ".docx")
    return (list)

def get_access_token():
  response = requests.post(
    "https://api.intra.42.fr/oauth/token",
    data={"grant_type": "client_credentials"},
    auth=("u-s4t2ud-xxxx", "s-s4t2ud-xxxx"),
  )
  return response.json()["access_token"]
